"""
Document loader supporting PDF, TXT, HTML, and DOCX formats.

Returns a normalised Document dataclass regardless of source format so that
downstream preprocessing and chunking stages have a consistent interface.
"""
from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path
from typing import Any


@dataclass
class Document:
    content: str
    metadata: dict[str, Any]
    source_id: str


def load_from_bytes(content: bytes, filename: str, source_id: str) -> Document:
    """Dispatch to the appropriate loader based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _load_pdf(content, filename, source_id)
    elif ext == ".docx":
        return _load_docx(content, filename, source_id)
    elif ext in (".htm", ".html"):
        return _load_html(content, filename, source_id)
    else:
        return _load_text(content, filename, source_id)


def load_from_file(path: str | Path, source_id: str | None = None) -> Document:
    """Load a document from a filesystem path."""
    p = Path(path)
    sid = source_id or p.stem
    with open(p, "rb") as f:
        raw = f.read()
    return load_from_bytes(raw, p.name, sid)


def _load_pdf(content: bytes, filename: str, source_id: str) -> Document:
    from pypdf import PdfReader  # type: ignore[import]

    reader = PdfReader(io.BytesIO(content))
    text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
    return Document(
        content=text,
        metadata={"filename": filename, "format": "pdf", "pages": len(reader.pages)},
        source_id=source_id,
    )


def _load_docx(content: bytes, filename: str, source_id: str) -> Document:
    import docx  # type: ignore[import]

    doc = docx.Document(io.BytesIO(content))
    text = "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
    return Document(
        content=text,
        metadata={"filename": filename, "format": "docx"},
        source_id=source_id,
    )


def _load_html(content: bytes, filename: str, source_id: str) -> Document:
    from bs4 import BeautifulSoup  # type: ignore[import]

    soup = BeautifulSoup(content, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    text = soup.get_text(separator="\n").strip()
    return Document(
        content=text,
        metadata={"filename": filename, "format": "html"},
        source_id=source_id,
    )


def _load_text(content: bytes, filename: str, source_id: str) -> Document:
    text = content.decode("utf-8", errors="replace")
    return Document(
        content=text,
        metadata={"filename": filename, "format": "txt"},
        source_id=source_id,
    )
